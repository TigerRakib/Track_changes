import os
import zipfile
import shutil
from lxml import etree
from deep_translator import GoogleTranslator
from docx import Document # Used for creating dummy docx if needed
import docx.oxml.ns # Used for setting properties in dummy docx

def modify_docx_preserve_format(input_docx, output_docx, text_to_add, target_language='zh-CN'):
    """
    Modifies a DOCX file by:
    1. Translating all existing text content to the target language.
    2. Removing specific runs (blue and strikethrough text).
    3. Adding a new paragraph with blue underlined text, also translated.
    The original formatting of the document is preserved as much as possible.

    Args:
        input_docx (str): Path to the input DOCX file.
        output_docx (str): Path to save the modified DOCX file.
        text_to_add (str): The text string to be added as a new paragraph.
        target_language (str): The language code for translation (e.g., 'zh-CN' for Chinese).
    """
    temp_dir = "temp_docx_extracted"

    # Ensure the temporary directory is clean before starting
    if os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)
    os.makedirs(temp_dir)

    print(f"Starting full DOCX translation and modification for '{input_docx}' to '{output_docx}'")
    print(f"Text to add (will also be translated): '{text_to_add}'")
    print(f"Target translation language: '{target_language}'")

    try:
        # 1. Unzip the .docx into a temporary folder
        with zipfile.ZipFile(input_docx, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        print(f"Extracted DOCX to '{temp_dir}'")

        document_path = os.path.join(temp_dir, "word", "document.xml")

        # Check if document.xml exists
        if not os.path.exists(document_path):
            print(f"Error: document.xml not found in {input_docx}. Is it a valid DOCX file?")
            return

        # 2. Parse the document.xml
        parser = etree.XMLParser(remove_blank_text=False)
        tree = etree.parse(document_path, parser)
        root = tree.getroot()
        print(f"Parsed document.xml from '{document_path}'")

        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

        translator = GoogleTranslator(source='auto', target=target_language)

        # 3. Translate all existing text content in the document
        print("Translating existing document text...")
        text_elements = root.xpath('//w:t', namespaces=ns)
        translated_count = 0
        total_text_elements = len(text_elements)

        for i, text_elem in enumerate(text_elements):
            original_text = text_elem.text
            if original_text and original_text.strip(): # Only translate non-empty text
                # Print progress before attempting translation
                print(f"  Attempting to translate segment {i+1}/{total_text_elements}: '{original_text[:50]}...'")
                try:
                    translated_segment = translator.translate(original_text)
                    text_elem.text = translated_segment
                    translated_count += 1
                    print(f"  Successfully translated segment {i+1}.")
                except Exception as e:
                    # Log specific error type and message if translation fails
                    print(f"  WARNING: Translation failed for segment {i+1} ('{original_text[:50]}...'). Error: {type(e).__name__}: {e}. Keeping original text.")
            elif original_text is None: # Handle cases where text_elem.text might be None
                text_elem.text = "" # Ensure it's an empty string if it was None
        print(f"Finished translating {translated_count} text segments in the document.")

        # 4. Remove runs that are blue and strikethrough
        runs_to_remove = []
        for run in root.xpath('//w:r', namespaces=ns):
            rpr = run.find('w:rPr', namespaces=ns)
            text = run.find('w:t', namespaces=ns)
            if rpr is not None and text is not None:
                color = rpr.find('w:color', namespaces=ns)
                strike = rpr.find('w:strike', namespaces=ns)
                if (color is not None and color.get(w + 'val', '').lower() == '0000ff' and
                    strike is not None and strike.get(w + 'val', 'true') != 'false'):
                    runs_to_remove.append(run)

        for run in runs_to_remove:
            parent = run.getparent()
            if parent is not None:
                parent.remove(run)
        print(f"Removed {len(runs_to_remove)} blue strikethrough runs.")

        # 5. Translate the text_to_add for the new paragraph
        translated_text_to_add = text_to_add # Initialize with original text as fallback
        try:
            translated_text_to_add = translator.translate(text_to_add)
            print(f"Original text for new paragraph: '{text_to_add}'")
            print(f"Translated text for new paragraph ({target_language}): '{translated_text_to_add}'")
        except Exception as e:
            print(f"ERROR: Translation for new paragraph failed! Reason: {type(e).__name__}: {e}. Falling back to original text for new paragraph insertion.")

        # 6. Add a new paragraph with blue underlined text
        body = root.find('w:body', namespaces=ns)
        if body is None:
            print("Error: Could not find 'w:body' element in document.xml.")
            return

        p = etree.Element(w + "p")
        r = etree.SubElement(p, w + "r")
        rpr = etree.SubElement(r, w + "rPr")

        color = etree.SubElement(rpr, w + "color")
        color.set(w + "val", "0000FF") # Blue color

        underline = etree.SubElement(rpr, w + "u")
        underline.set(w + "val", "single") # Single underline

        t = etree.SubElement(r, w + "t")
        t.text = translated_text_to_add # Use the translated text for the new paragraph
        body.append(p)
        print(f"Added new paragraph with text: '{translated_text_to_add}'")

        # 7. Save the modified XML
        tree.write(document_path, xml_declaration=True, encoding='UTF-8', standalone='yes')
        print(f"Saved modified document.xml to '{document_path}'")

        # 8. Zip everything back into a new .docx
        output_zip_path = "new_docx.zip"
        with zipfile.ZipFile(output_zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root_dir, _, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root_dir, file)
                    arcname = os.path.relpath(file_path, temp_dir)
                    zipf.write(file_path, arcname)
        print(f"Zipped contents back to '{output_zip_path}'")

        shutil.move(output_zip_path, output_docx)
        print(f"Successfully created '{output_docx}'.")

    except Exception as e:
        print(f"An unexpected error occurred during DOCX processing: {e}")
    finally:
        # 9. Clean up
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
            print(f"Cleaned up temporary directory: {temp_dir}")

# === Example usage ===
input_file = "[Track Changes] KFS - AXA Global Strategic Bonds_E.docx" # Make sure you have a 'demo.docx' file in the same directory
output_file = "output_full_chinese.docx" # Changed output filename to reflect full translation
text_to_add = "This is a new paragraph added to the document."
target_language = "zh-CN" # Chinese (Simplified)

# Create a dummy demo.docx for testing if it doesn't exist
if not os.path.exists(input_file):
    doc = Document()
    doc.add_paragraph("Hello, this is the first paragraph in English.")
    doc.add_paragraph("This is the second paragraph with some numbers: 123.")
    # Add a blue strikethrough run for testing removal
    p = doc.add_paragraph()
    run = p.add_run("This text should be removed and is blue and strikethrough.")
    rpr = run._r.get_or_add_rPr()
    color_elm = etree.SubElement(rpr, docx.oxml.ns.qn('w:color'))
    color_elm.set(docx.oxml.ns.qn('w:val'), '0000FF')
    strike_elm = etree.SubElement(rpr, docx.oxml.ns.qn('w:strike'))
    strike_elm.set(docx.oxml.ns.qn('w:val'), 'true')

    doc.add_paragraph("Another paragraph here to be translated.")
    doc.save(input_file)
    print(f"Created a dummy '{input_file}' for demonstration.")

modify_docx_preserve_format(input_file, output_file, text_to_add, target_language)
